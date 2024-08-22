from docx import Document

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK

from docx.oxml.ns import qn

from docx.oxml import OxmlElement

from docx.shared import Cm, Pt


import locale

from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.section import _Header
from docx.section import _Footer

from typing import List


from tkinter import filedialog

from datetime import datetime

import os

import subprocess


class PropostaSocietaria:

    def __init__(self, nome_empresa:str, mes_ano:str, codigo:str, data:str, nome_responsavel:str,
                 proposta:str, lista_termos_proposta:List, valor_honorario:str) -> None:
        
        self.code_keys:dict ={
            'StrNomeEmpresa':(nome_empresa, 'bold'),

            'StrMesAno': (mes_ano, 'default'),

            'StrCodigo': (codigo, 'bold'),

            'StrData':(data, 'default'),

            'StrNomeResponsavel': (nome_responsavel, 'bold'),

            'StrProposta': (proposta, 'bold'),

            'ListBulletTermosProposta': (lista_termos_proposta, 'ListBullet'),

            'StrValorHonorario': (valor_honorario, 'bold')

        }

        self.documento = Document('src/modelos/modelo_proposta_societaria.docx')

        style = self.documento.styles['Normal']

        font = style.font

        font.name = 'Arial Narrow'
        font.size = Pt(12)

    def gerar_proposta_societaria(self, path:str):

        for paragrafo in self.documento.paragraphs:

            for key in list(self.code_keys.keys()):

                if key in paragrafo.text and key != 'ListBulletTermosProposta' :


                    texto_original = paragrafo.text

                    paragrafo.clear()

                    partes = texto_original.split(key)

                    for i, parte in enumerate(partes):

                        if i>0:
                            run = paragrafo.add_run(self.code_keys[key][0])

                            if self.code_keys[key][1] == 'bold':

                                run.bold = True

                        paragrafo.add_run(parte)

                    
                if key in paragrafo.text and key == 'ListBulletTermosProposta':

                    paragrafo.clear()

                    for item in self.code_keys[key][0]:
                        paragrafo.insert_paragraph_before(text = f'\t\t• {item} ')

        full_path = f'{path}/{self.code_keys['StrCodigo'][0]}_proposta_societaria_{(self.code_keys['StrNomeEmpresa'][0]).replace(' ', '_')}_{datetime.today().month}_{datetime.today().year}'

        self.documento.save(f'{full_path}.docx')

        subprocess.run(f'start {full_path}.docx', shell = True)

